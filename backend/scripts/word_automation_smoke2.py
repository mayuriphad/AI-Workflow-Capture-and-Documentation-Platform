"""Smoke test for the new word_automation methods: insert_step_after,
delete_step, replace_step_image -- through the real async queue bridge."""
import asyncio
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.services.word_automation import WordAutomation  # noqa: E402

STORAGE = Path(__file__).resolve().parent.parent / "storage"
SAMPLE_IMAGE = next(STORAGE.glob("screenshots/*.png"), None)
DOCX_PATH = STORAGE / "word_automation_smoke2.docx"


async def main() -> None:
    word = WordAutomation()
    word.start()
    try:
        await word.create_document("proj-smoke2", "Smoke Test 2", str(DOCX_PATH))
        await word.append_step("Step one instruction", str(SAMPLE_IMAGE), 1, "step_1")
        await word.append_step("Step two instruction", str(SAMPLE_IMAGE), 2, "step_2")
        print("appended 2 steps")

        await word.insert_step_after("step_1", "Inserted between 1 and 2", str(SAMPLE_IMAGE), 99, "step_mid")
        print("inserted step between step_1 and step_2")

        await word.replace_step_image("step_mid", str(SAMPLE_IMAGE))
        print("replaced image on step_mid")

        await word.delete_step("step_1")
        print("deleted step_1")

        from docx import Document as DocxDocument

        doc = DocxDocument(str(DOCX_PATH))
        text = "\n".join(p.text for p in doc.paragraphs)
        has_step1 = "Step one instruction" in text
        has_mid = "Inserted between 1 and 2" in text
        has_step2 = "Step two instruction" in text
        order_ok = text.index("Inserted between 1 and 2") < text.index("Step two instruction")
        print(f"step_1 gone (expect False): {has_step1}")
        print(f"mid step present: {has_mid}")
        print(f"step2 present: {has_step2}")
        print(f"order ok: {order_ok}")

        await word.delete_step("step_mid")
        await word.delete_step("step_2")
        doc2 = DocxDocument(str(DOCX_PATH))
        remaining = "\n".join(p.text for p in doc2.paragraphs).strip()
        print(f"doc empty-ish after deleting all steps: {remaining!r}")

        ok = (not has_step1) and has_mid and has_step2 and order_ok
        print("RESULT:", "PASS" if ok else "FAIL")

        await word.close(save=False)
    finally:
        word.stop()


if __name__ == "__main__":
    asyncio.run(main())
